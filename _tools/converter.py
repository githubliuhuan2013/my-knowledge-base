"""Word (.docx) → Markdown 转换器

使用 python-docx 解析 Word 文档，转为 Markdown 格式。
处理：标题、段落、粗体/斜体、列表、表格、图片占位。
"""

import os
import re
from pathlib import Path
from docx import Document


def docx_to_markdown(docx_path: str) -> str:
    """将 .docx 文件转换为 Markdown 文本"""
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = _extract_paragraph_text(para)

        if not text.strip():
            lines.append("")
            continue

        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            level_match = re.search(r"\d+", style_name)
            level = int(level_match.group()) if level_match else 1
            level = min(level, 6)
            lines.append(f"{'#' * level} {text}")
            lines.append("")

        elif _is_list_item(para, style_name):
            lines.append(f"- {text}")

        elif _is_blockquote(para, style_name):
            lines.append(f"> {text}")

        else:
            lines.append(text)
            lines.append("")

    # 处理表格
    for table in doc.tables:
        lines.append("")
        lines.append(_table_to_markdown(table))
        lines.append("")

    return "\n".join(lines)


def _extract_paragraph_text(para) -> str:
    """提取段落文本，保留格式标记"""
    result = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            t = f"***{t}***"
        elif run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        elif run.underline:
            t = f"<u>{t}</u>"
        result.append(t)
    return "".join(result)


def _is_list_item(para, style_name: str) -> bool:
    """判断是否为列表项"""
    list_keywords = ["List", "Bullet", "Number", "Bulleted", "Numbered"]
    return any(kw in style_name for kw in list_keywords)


def _is_blockquote(para, style_name: str) -> bool:
    return "Quote" in style_name or "Block" in style_name


def _table_to_markdown(table) -> str:
    """将 Word 表格转为 Markdown 表格"""
    rows = table.rows
    if not rows:
        return ""

    result = []
    for i, row in enumerate(rows):
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        result.append("| " + " | ".join(cells) + " |")
        if i == 0:
            result.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(result)


def extract_title_from_docx(docx_path: str) -> str:
    """从 docx 文件名或内容推测标题"""
    filename = os.path.splitext(os.path.basename(docx_path))[0]
    return filename


def convert_and_save(docx_path: str, output_dir: str) -> tuple[str, str]:
    """
    转换 docx 为 md 并保存。
    返回 (md_path, title)
    """
    md_content = docx_to_markdown(docx_path)
    title = extract_title_from_docx(docx_path)

    # 生成安全的文件名
    safe_name = re.sub(r'[<>:"/\\|?*]', "-", title)
    md_path = os.path.join(output_dir, f"{safe_name}.md")

    # 添加 frontmatter
    from datetime import date
    frontmatter = f"""---
title: "{title}"
tags: []
created: {date.today().isoformat()}
updated: {date.today().isoformat()}
links: []
source: "_originals/{os.path.basename(docx_path)}"
---

"""
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(frontmatter + md_content)

    return md_path, title
