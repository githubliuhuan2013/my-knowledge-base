"""生成知识库方案 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True

# === 封面 ===
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('个人智能知识库体系\n建设方案')
run.font.size = Pt(28)
run.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于 GitHub + Markdown + AI 的个人知识管理方案')
run.font.size = Pt(14)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('创建日期：2026年7月6日\n工具：Claude Code + Python + GitHub')
run.font.size = Pt(11)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# === 目录页 ===
add_heading_styled('目录', 1)
toc_items = [
    '一、项目背景与目标',
    '二、整体架构设计',
    '三、知识库目录结构',
    '四、Phase 1：知识库骨架',
    '五、Phase 2：智能搜索',
    '六、Phase 3：知识关联图谱',
    '七、Phase 4：Web 界面（可选）',
    '八、技术栈总览',
    '九、使用指南',
    '十、迭代路线图',
]
for item in toc_items:
    add_para(item)

doc.add_page_break()

# === 一、项目背景与目标 ===
add_heading_styled('一、项目背景与目标', 1)
add_para('在日常工作和学习中，我们会积累大量的文档、笔记、资料。传统的文件管理方式（一个个文件夹 + Ctrl+F 搜索）效率低下，无法体现知识之间的关联，也难以跨设备访问。')
add_para('')
add_para('本方案旨在打造一个智能个人知识库，具备以下核心能力：')
add_para('✅ 资料统一管理：支持 Markdown 笔记 + Word 文档混合存储', bold=True)
add_para('✅ 跨设备同步：基于 GitHub 私有仓库，办公室和家里电脑自动同步', bold=True)
add_para('✅ 聊天式检索：不再 Ctrl+F，而是用自然语言对话查找知识', bold=True)
add_para('✅ 知识关联：自动发现笔记之间的联系，构建知识图谱', bold=True)
add_para('✅ 持续迭代：分阶段建设，每阶段可独立使用并产生价值', bold=True)

# === 二、整体架构设计 ===
add_heading_styled('二、整体架构设计', 1)
add_para('整个系统以 GitHub 私有仓库为核心中枢，本地运行 Python 智能工具，Claude Code 提供 AI 聊天检索入口。')
add_para('')

# 架构图
add_para('┌──────────────────────────────────────────┐')
add_para('│           GitHub 私有仓库（中枢）          │')
add_para('│  ┌──────────┐  ┌──────────┐  ┌─────────┐ │')
add_para('│  │ Markdown │  │ 原始Word │  │ Actions │ │')
add_para('│  │  知识库   │  │  文档归档 │  │ 自动处理 │ │')
add_para('│  └──────────┘  └──────────┘  └─────────┘ │')
add_para('└──────────────────────────────────────────┘')
add_para('         ↕ git push/pull（跨设备同步）')
add_para('┌──────────────────┐    ┌──────────────────┐')
add_para('│    办公室电脑     │    │     家里电脑      │')
add_para('│  clone 仓库      │    │   clone 仓库      │')
add_para('│  kb CLI + Claude │    │  kb CLI + Claude  │')
add_para('└──────────────────┘    └──────────────────┘')

add_para('')
add_para('各层技术选型：')
add_para('• 存储+同步：GitHub 私有仓库（免费、跨设备、版本控制）')
add_para('• 本地工具：Python CLI 工具 kb（文档转换、索引、关联）')
add_para('• 向量搜索：ChromaDB + sentence-transformers（本地运行、免费）')
add_para('• AI 对话：Claude Code（已安装，自然语言理解 + 遍历检索）')
add_para('• 文档转换：python-docx（Word → Markdown）')

# === 三、知识库目录结构 ===
add_heading_styled('三、知识库目录结构', 1)
add_para('knowledge-base/                   # GitHub 仓库根目录')
add_para('├── notes/                        # 所有知识笔记（Markdown）')
add_para('│   ├── 技术/                     #   技术类笔记')
add_para('│   ├── 产品/                     #   产品类笔记')
add_para('│   ├── 阅读/                     #   读书笔记')
add_para('│   └── 其他/                     #   其他分类')
add_para('├── _originals/                   # 原始 Word 文档归档')
add_para('├── _templates/                   # 笔记模板')
add_para('│   └── note-template.md')
add_para('├── _tools/                       # Python 工具脚本')
add_para('│   ├── kb.py                     #   CLI 主入口')
add_para('│   ├── converter.py              #   Word → MD 转换')
add_para('│   └── indexer.py                #   向量索引构建')
add_para('├── _index/                       # ChromaDB 索引数据（gitignore）')
add_para('├── .github/workflows/            # GitHub Actions 自动化')
add_para('└── README.md                     # 仓库说明')

# === 四、Phase 1 ===
add_heading_styled('四、Phase 1：知识库骨架', 1)
add_para('目标：搭好仓库、能创建笔记、能导入 Word 文档。', bold=True)
add_para('')
add_para('4.1 笔记格式（YAML Frontmatter）')
add_para('每篇笔记使用标准 Markdown + YAML 头部元数据：')
add_para('---')
add_para('title: "Docker 常用命令"')
add_para('tags: [docker, devops, 运维]')
add_para('created: 2026-07-06')
add_para('updated: 2026-07-06')
add_para('links:')
add_para('  - "../技术/K8s入门.md"')
add_para('source: "_originals/Docker笔记.docx"')
add_para('---')
add_para('')
add_para('4.2 CLI 命令')
add_para('• kb new "标题" --tags tag1,tag2      → 创建新笔记')
add_para('• kb import ~/Desktop/资料.docx        → 导入 Word，自动转 MD')
add_para('• kb list --tag docker                 → 按标签列出笔记')
add_para('• kb status                            → 查看仓库状态')

# === 五、Phase 2 ===
add_heading_styled('五、Phase 2：智能搜索', 1)
add_para('目标：用自然语言提问，系统自动找到最相关的内容并合成回答。', bold=True)
add_para('')
add_para('5.1 工作原理（RAG 检索增强生成）')
add_para('用户问题 → 向量化 → ChromaDB 语义检索 Top-K 段落')
add_para('→ 拼接上下文 → Claude AI 合成回答（附来源引用）')
add_para('')
add_para('5.2 两种搜索方式')
add_para('• 方式A（日常快速查）：在 Claude Code 中直接对话，"在我的知识库里关于 XXX 有哪些内容？"')
add_para('• 方式B（精准语义搜）：kb ask "关于 XXX 有哪些内容？" → 深度语义搜索 + AI 合成')
add_para('')
add_para('5.3 技术组件')
add_para('• ChromaDB：轻量级向量数据库，文件存储，无需服务进程')
add_para('• sentence-transformers：本地 embedding 模型，免费，离线运行')
add_para('• Claude Code / API：AI 对话与内容合成')

# === 六、Phase 3 ===
add_heading_styled('六、Phase 3：知识关联图谱', 1)
add_para('目标：让知识之间自动建立连接，形成网状知识结构。', bold=True)
add_para('')
add_para('6.1 功能')
add_para('• 自动发现：基于内容相似度，自动发现相关笔记')
add_para('• 双向链接：解析笔记间的 links 字段，构建有向图')
add_para('• 关联推荐：kb relate → 扫描全库，列出建议关联')
add_para('• 图谱可视化：kb graph → 生成交互式知识图谱')
add_para('')
add_para('6.2 命令')
add_para('• kb links "笔记.md"    → 查看正向链接 + 反向链接 + 推荐链接')
add_para('• kb relate             → 扫描全库关联建议')
add_para('• kb graph              → 生成知识图谱')

# === 七、Phase 4 ===
add_heading_styled('七、Phase 4：Web 界面（可选）', 1)
add_para('目标：提供浏览器访问的图形化界面，方便浏览和上传。', bold=True)
add_para('')
add_para('7.1 功能模块')
add_para('• 笔记浏览（按目录/标签筛选）')
add_para('• 聊天搜索框（自然语言提问）')
add_para('• 拖拽上传 Word 文档')
add_para('• 知识图谱可视化')
add_para('• 下载原始 Word 文档')
add_para('')
add_para('7.2 技术选型')
add_para('• Streamlit（Python，极简 Web 框架）')
add_para('• 本地运行：kb web → 浏览器访问 localhost:8501')
add_para('• 需要远程访问时可部署到云服务器')

# === 八、技术栈总览 ===
add_heading_styled('八、技术栈总览', 1)

table = doc.add_table(rows=10, cols=3, style='Light Grid Accent 1')
headers = ['组件', '选型', '理由']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['版本控制+同步', 'Git + GitHub 私有仓库', '免费、跨设备、有 Actions'],
    ['文档解析', 'python-docx', '纯 Python，Windows 友好'],
    ['Embedding', 'sentence-transformers', '本地运行，免费，80MB'],
    ['向量数据库', 'ChromaDB', '轻量，文件存储，无需服务'],
    ['CLI 框架', 'click', 'Python CLI 标准库'],
    ['Web UI（可选）', 'Streamlit', '极简，适合个人工具'],
    ['AI 对话', 'Claude Code + API', '高质量中文理解'],
    ['Word 生成', 'python-docx', '生成方案文档'],
    ['跨设备同步', 'git push/pull', '通过 GitHub 中转'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# === 九、使用指南 ===
add_heading_styled('九、使用指南', 1)

add_para('9.1 日常使用流程', bold=True)
add_para('① 整理资料时：kb import ~/Desktop/某文档.docx')
add_para('② 记录笔记时：kb new "笔记标题" --tags 标签')
add_para('③ 查找知识时：打开 Claude Code 直接对话提问')
add_para('④ 下班前：kb sync（自动 git add + commit + push）')
add_para('⑤ 另一台电脑：kb sync（自动 git pull 同步最新内容）')
add_para('')
add_para('9.2 在 Claude Code 中搜索知识库', bold=True)
add_para('直接对 Claude Code 说：')
add_para('"在我的知识库 E:/Rex/knowledge-base/notes/ 中搜索关于 XXX 的所有内容"')
add_para('（Claude Code 会自动遍历所有 Markdown 文件找到相关内容）')
add_para('')
add_para('9.3 跨设备使用', bold=True)
add_para('办公室电脑和家里电脑各 clone 一次仓库，之后只需 kb sync 同步即可。')

# === 十、迭代路线图 ===
add_heading_styled('十、迭代路线图', 1)
add_para('Phase 1 ████████░░  基础骨架     已完成：仓库、目录、CLI、Word导入转换')
add_para('Phase 2 ████████░░  智能搜索     即将：向量索引、RAG、聊天查询')
add_para('Phase 3 ████████░░  知识图谱     计划：自动关联、双向链接、可视化')
add_para('Phase 4 ████████░░  Web界面      可选：Streamlit、拖拽上传')
add_para('未来    ░░░░░░░░░░  持续迭代     OCR、浏览器剪藏、移动端...')

doc.add_paragraph()
doc.add_paragraph()
add_para('---')
add_para('本文档由 Claude Code 自动生成，方案实施同步进行中。', bold=True)

# 保存
output_path = 'E:/Rex/个人智能知识库体系建设方案.docx'
doc.save(output_path)
print(f'[OK] Word document generated: {output_path}')
print(f'     Size: {os.path.getsize(output_path):,} bytes')
